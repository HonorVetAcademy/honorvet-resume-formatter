import re

from docx.shared import Pt

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _add_run_with_bold_markers(paragraph, text: str, base_bold=False, size=10.5):
    """Add text to a paragraph, honoring **bold** markers left over from AI-authored summary bullets."""
    pos = 0
    for match in BOLD_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.bold = base_bold
            run.font.size = Pt(size)
        run = paragraph.add_run(match.group(1))
        run.bold = True
        run.font.size = Pt(size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.font.size = Pt(size)


def _section_header(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    return p


def _bullet(doc, text: str, indent=False):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Pt(28)
    p.paragraph_format.space_after = Pt(3)
    _add_run_with_bold_markers(p, text)
    return p


def _plain_line(doc, label: str, value):
    if value in (None, "", "null"):
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{label}: ")
    run.font.size = Pt(10.5)
    run2 = p.add_run(str(value))
    run2.font.size = Pt(10.5)
