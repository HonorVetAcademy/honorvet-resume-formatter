import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.resume_docx_generator import _section_header, _bullet, _plain_line


def generate_rightsourcing_docx(resume: dict, output_dir: str) -> str:
    """Render a structured resume into the HonorVet standard submission format:
    intro table, summary, skills, education, licenses & certs, professional experience."""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(60)
        section.right_margin = Pt(60)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Header
    name_line = resume.get("full_name", "")
    if resume.get("credentials_suffix"):
        name_line += f", {resume['credentials_suffix']}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name_line)
    run.bold = True
    run.font.size = Pt(15)

    # Introduction table
    _section_header(doc, "Introduction:")
    intro_rows = [
        ("Phone", resume.get("phone", "")),
        ("Email", resume.get("email", "")),
        ("Permanent Address", resume.get("permanent_address", "") or "— fill in —"),
        ("Available to Start Date", "— fill in (mm/dd/yyyy) —"),
        ("Weekends/Holiday Availability", "— fill in (y/n) —"),
    ]
    table = doc.add_table(rows=len(intro_rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(intro_rows):
        cell_label = table.rows[i].cells[0]
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(10)
        cell_value = table.rows[i].cells[1]
        cell_value.text = str(value)
        cell_value.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Professional Summary
    if resume.get("professional_summary"):
        _section_header(doc, "Professional Summary:")
        for bullet in resume["professional_summary"]:
            _bullet(doc, bullet)

    # Skills
    if resume.get("skills"):
        _section_header(doc, "Skills:")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(", ".join(resume["skills"]))
        run.font.size = Pt(10.5)

    # Education
    if resume.get("education"):
        _section_header(doc, "Education:")
        for edu in resume["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(edu.get("degree", ""))
            run.bold = True
            run.font.size = Pt(10.5)
            tail = " ".join(x for x in [edu.get("school", ""), "–", edu.get("location", "")] if x)
            r2 = p.add_run(f" {tail}")
            r2.font.size = Pt(10.5)
            if edu.get("date"):
                r3 = p.add_run(f" | {edu['date']}")
                r3.bold = True
                r3.font.size = Pt(10.5)

    # Licenses and Certifications
    if resume.get("certifications"):
        _section_header(doc, "Licenses and Certifications:")
        for cert in resume["certifications"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            label = cert.get("name", "")
            if cert.get("issuer"):
                label += f" – {cert['issuer']}"
            run = p.add_run(label)
            run.font.size = Pt(10.5)
            extras = []
            if cert.get("id"):
                extras.append(f"# {cert['id']}")
            if cert.get("expires"):
                extras.append(f"Expires: {cert['expires']}")
            if extras:
                r2 = p.add_run(" | " + " | ".join(extras))
                r2.bold = True
                r2.font.size = Pt(10.5)

    # Professional Experience
    if resume.get("experience"):
        _section_header(doc, "Professional Experience:")
        for job in resume["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            loc = ", ".join(x for x in [job.get("city", ""), job.get("state", "")] if x)
            facility_line = job.get("facility_name", "")
            if loc:
                facility_line += f" | {loc}"
            run = p.add_run(facility_line)
            run.bold = True
            run.font.size = Pt(10.5)
            dates = f"{job.get('start_date', '')} to {job.get('end_date', '')}"
            r2 = p.add_run(f" | {dates}")
            r2.bold = True
            r2.font.size = Pt(10.5)

            if job.get("job_title"):
                tp = doc.add_paragraph()
                tp.paragraph_format.space_after = Pt(2)
                tr = tp.add_run(job["job_title"])
                tr.bold = True
                tr.font.size = Pt(10.5)

            _plain_line(doc, "EMR", job.get("emr_system"))
            _plain_line(doc, "Position Type", job.get("position_type") or "— fill in —")
            _plain_line(doc, "Agency Name", job.get("agency_name") or "— fill in —")
            _plain_line(doc, "Trauma Level", job.get("trauma_level"))
            _plain_line(doc, "Facility Type", job.get("facility_type"))

            for duty in job.get("duties", []):
                _bullet(doc, duty, indent=True)

    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", resume.get("full_name", "candidate"))
    filename = f"{safe_name}_HonorVet_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath
